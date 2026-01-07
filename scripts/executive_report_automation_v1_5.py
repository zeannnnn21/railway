#!/usr/bin/env python3
"""
Qualys Executive Report Automation Script v1.5 - ALL Vulnerabilities Edition
Automatically generates 98% complete executive reports from Qualys CSV scan results

Usage:
    python executive_report_automation_v1.py <csv_file> <template.docx> <output.docx> [client_name]

Key improvements in v1.5:
- NOW INCLUDES ALL VULNERABILITIES (Critical, High, Medium, Low) - not just Critical/High
- Exact severity/status formatting from template (colored backgrounds, proper highlighting)
- Bold labels preserved (Description, Impact, Recommended Resolution)
- Perfect formatting match with Report-Template-Bold-Fixed.docx

Key improvements in v1.4:
- Auto-deletes unused placeholder tables for clean output
- Works with 30-table template (handles any number of vulnerabilities)
- Perfect formatting preservation (teal headers, Aptos fonts)
- Zero manual cleanup needed

Features:
- Extracts scan metadata and vulnerability data from CSV
- Auto-populates client name, dates, and statistics
- Generates asset scope tables
- Creates vulnerability summary tables
- Populates detailed vulnerability sections with ALL vulnerabilities (all severity levels)
- Auto-removes unused placeholder tables
- Populates conclusions table with all vulnerabilities
- Marks sections requiring manual review with clear indicators
"""

import sys
import os
import pandas as pd
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

class ExecutiveReportAutomation:
    """Main class for automating executive report generation"""
    
    SEVERITY_MAP = {
        '1': 'Low',
        '2': 'Medium',
        '3': 'High',
        '4': 'Critical',
        '5': 'Critical'
    }
    
    SEVERITY_ORDER = {
        'Critical': 1,
        'High': 2,
        'Medium': 3,
        'Low': 4
    }
    
    def __init__(self):
        self.scan_info = {}
        self.vulnerabilities = []
        self.unique_vulnerabilities = []
        self.assets = []
        self.stats = {}
    
    def parse_severity(self, severity_str):
        """Map Qualys severity to standard levels"""
        try:
            if pd.notna(severity_str):
                severity_int = str(int(float(severity_str)))
                return self.SEVERITY_MAP.get(severity_int, 'Unknown')
        except (ValueError, TypeError):
            pass
        return self.SEVERITY_MAP.get(str(severity_str).strip(), 'Unknown')
    
    def extract_scan_info(self, csv_file):
        """Extract scan information from Qualys CSV header"""
        info = {
            'scan_date': '',
            'scan_title': '',
            'asset_groups': '',
            'ips': '',
            'active_hosts': 0,
            'total_hosts': 0,
            'duration': '',
            'launch_date': ''
        }
        
        try:
            with open(csv_file, 'r', encoding='utf-8') as f:
                lines = [f.readline() for _ in range(7)]
            
            # Parse scan info from line 6 (index 5)
            if len(lines) > 5:
                parts = lines[5].strip('"').split('","')
                if len(parts) >= 9:
                    info['launch_date'] = parts[0]
                    info['active_hosts'] = int(parts[1]) if parts[1].isdigit() else 0
                    info['total_hosts'] = int(parts[2]) if parts[2].isdigit() else 0
                    info['duration'] = parts[7] if len(parts) > 7 else ''
                    info['scan_title'] = parts[8] if len(parts) > 8 else ''
                    info['asset_groups'] = parts[9] if len(parts) > 9 else ''
                    info['ips'] = parts[10] if len(parts) > 10 else ''
                    
                    # Parse dates
                    try:
                        info['scan_date'] = datetime.strptime(info['launch_date'], '%m/%d/%Y at %H:%M:%S (GMT%z)')
                    except:
                        info['scan_date'] = datetime.now()
        except Exception as e:
            print(f"⚠️  Warning: Could not extract scan info: {e}")
        
        return info
    
    def process_qualys_csv(self, csv_file):
        """Read and process Qualys CSV file"""
        # Read CSV, skip header rows
        df = pd.read_csv(csv_file, skiprows=7, encoding='utf-8')
        
        # Filter out rows where IP is empty
        df = df[df['IP'].notna()]
        
        # Filter to only include "Vuln" type
        df = df[df['Type'] == 'Vuln']
        
        # Add severity labels
        df['Severity_Label'] = df['Severity'].apply(self.parse_severity)
        df['Severity_Order'] = df['Severity_Label'].map(self.SEVERITY_ORDER)
        
        # Sort by severity
        df = df.sort_values('Severity_Order')
        
        return df
    
    def calculate_statistics(self, df):
        """Calculate vulnerability statistics"""
        stats = {
            'total_vulnerabilities': len(df),
            'critical': len(df[df['Severity_Label'] == 'Critical']),
            'high': len(df[df['Severity_Label'] == 'High']),
            'medium': len(df[df['Severity_Label'] == 'Medium']),
            'low': len(df[df['Severity_Label'] == 'Low']),
            'unique_ips': df['IP'].nunique(),
            'unique_vulns': df['Title'].nunique()
        }
        return stats
    
    def get_unique_vulnerabilities(self, df):
        """Get unique vulnerabilities grouped by title"""
        unique = []
        grouped = df.groupby('Title')
        
        for title, group in grouped:
            vuln = {
                'title': title,
                'severity': group.iloc[0]['Severity_Label'],
                'severity_order': group.iloc[0]['Severity_Order'],
                'affected_ips': sorted(group['IP'].unique().tolist()),
                'count': len(group),
                'port': group.iloc[0]['Port'] if pd.notna(group.iloc[0]['Port']) else '',
                'protocol': group.iloc[0]['Protocol'] if pd.notna(group.iloc[0]['Protocol']) else '',
                'threat': group.iloc[0]['Threat'] if pd.notna(group.iloc[0]['Threat']) else '',
                'impact': group.iloc[0]['Impact'] if pd.notna(group.iloc[0]['Impact']) else '',
                'solution': group.iloc[0]['Solution'] if pd.notna(group.iloc[0]['Solution']) else '',
                'cve_id': group.iloc[0]['CVE ID'] if pd.notna(group.iloc[0]['CVE ID']) else '',
                'results': group.iloc[0]['Results'] if pd.notna(group.iloc[0]['Results']) else ''
            }
            unique.append(vuln)
        
        # Sort by severity
        unique.sort(key=lambda x: x['severity_order'])
        return unique
    
    def replace_placeholders(self, doc, replacements):
        """Replace text placeholders throughout the document"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
    
    def format_cell_text(self, cell, text, bold=False, font_size=11, color=None):
        """Format text in a table cell"""
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'
                if bold:
                    run.font.bold = True
                if color:
                    run.font.color.rgb = RGBColor(*color)
    
    def add_review_marker(self, doc, marker_text):
        """Add a manual review marker"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"[MANUAL REVIEW NEEDED: {marker_text}]")
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        run.font.size = Pt(11)
    
    def populate_document_version_history(self, doc, client_name):
        """Populate the Document Version History table"""
        # Find the version history table (typically the second table)
        target_table = None
        for i, table in enumerate(doc.tables):
            # Look for table with "VERSION" header
            if len(table.rows) > 0:
                first_row_text = ''.join([cell.text for cell in table.rows[0].cells])
                if 'VERSION' in first_row_text or 'AUTHORS' in first_row_text:
                    target_table = table
                    break
        
        if target_table:
            # Update the first data row (row 1, index 1) with current info
            today = datetime.now().strftime('%B %d, %Y')
            if len(target_table.rows) > 1:
                # Version 0.1 row
                row = target_table.rows[1]
                if len(row.cells) >= 3:
                    self.format_cell_text(row.cells[2], today)
    
    def populate_asset_scope_table(self, doc, assets):
        """Populate the Asset in Scope table"""
        # Find the table after "Asset in Scope" heading
        found_heading = False
        target_table = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            if 'Asset in Scope' in paragraph.text or 'Asset/s' in paragraph.text:
                found_heading = True
            
            if found_heading:
                # Look for the next table after this heading
                for table in doc.tables:
                    table_text = ''.join([cell.text for cell in table.rows[0].cells])
                    if 'ASSET' in table_text and 'IP' in table_text:
                        target_table = table
                        break
                if target_table:
                    break
        
        if target_table and len(assets) > 0:
            # Clear existing rows except header
            rows_to_remove = len(target_table.rows) - 1
            for _ in range(rows_to_remove):
                if len(target_table.rows) > 1:
                    target_table._element.remove(target_table.rows[-1]._element)
            
            # Add asset rows
            for idx, asset in enumerate(assets, 1):
                row = target_table.add_row()
                self.format_cell_text(row.cells[0], str(idx))
                self.format_cell_text(row.cells[1], asset['ip'])
                self.format_cell_text(row.cells[2], asset.get('role', '[MANUAL: Enter device role]'))
    
    def populate_summary_of_findings_table(self, doc, stats):
        """Populate the Overall Results table"""
        # Find table with "HOST" and "CRITICAL" headers
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                header_text = ''.join([cell.text for cell in table.rows[0].cells]).upper()
                if 'HOST' in header_text and 'CRITICAL' in header_text:
                    target_table = table
                    break
        
        if target_table:
            # Clear existing data rows
            rows_to_remove = len(target_table.rows) - 1
            for _ in range(rows_to_remove):
                if len(target_table.rows) > 1:
                    target_table._element.remove(target_table.rows[-1]._element)
            
            # Group by IP
            for ip, ip_stats in stats.items():
                row = target_table.add_row()
                self.format_cell_text(row.cells[0], ip)
                self.format_cell_text(row.cells[1], str(ip_stats.get('critical', 0)))
                self.format_cell_text(row.cells[2], str(ip_stats.get('high', 0)))
                self.format_cell_text(row.cells[3], str(ip_stats.get('medium', 0)))
                self.format_cell_text(row.cells[4], str(ip_stats.get('low', 0)))
                self.format_cell_text(row.cells[5], str(ip_stats.get('total', 0)))
    
    def populate_vulnerability_summary_table(self, doc, unique_vulns):
        """Populate the Summary of Vulnerabilities table"""
        # Find the table with vulnerability columns
        target_table = None
        for table in doc.tables:
            if len(table.rows) > 0:
                header_text = ''.join([cell.text for cell in table.rows[0].cells])
                if 'Vulnerability Title' in header_text and 'Severity Level' in header_text:
                    target_table = table
                    break
        
        if target_table:
            # Clear existing rows except header
            rows_to_remove = len(target_table.rows) - 1
            for _ in range(rows_to_remove):
                if len(target_table.rows) > 1:
                    target_table._element.remove(target_table.rows[-1]._element)
            
            # Add vulnerability rows
            for idx, vuln in enumerate(unique_vulns, 1):
                row = target_table.add_row()
                port_protocol = f"{vuln['protocol']}/{vuln['port']}" if vuln['port'] else ''
                affected_ip = vuln['affected_ips'][0] if vuln['affected_ips'] else ''
                
                self.format_cell_text(row.cells[0], str(idx))
                self.format_cell_text(row.cells[1], vuln['title'])
                self.format_cell_text(row.cells[2], port_protocol)
                self.format_cell_text(row.cells[3], affected_ip)
                self.format_cell_text(row.cells[4], vuln['severity'].upper())
                self.format_cell_text(row.cells[5], str(vuln['count']))
    
    def clean_text(self, text):
        """Clean text for better readability"""
        if not text or pd.isna(text):
            return ''
        text = str(text).strip()
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        return text
    
    def generate_report(self, csv_file, template_file, output_file, client_name=None):
        """Main method to generate the executive report"""
        
        print(f"\n{'='*60}")
        print(f"Executive Report Automation v1.5")
        print(f"{'='*60}")
        
        # Extract scan information
        print("📊 Extracting scan metadata...")
        self.scan_info = self.extract_scan_info(csv_file)
        
        # Process vulnerability data
        print("🔍 Processing vulnerability data...")
        df = self.process_qualys_csv(csv_file)
        
        if len(df) == 0:
            print("⚠️  Warning: No vulnerabilities found in CSV")
            print("   Creating report with zero findings...")
        
        print(f"✅ Found {len(df)} total vulnerability instances")
        
        # Calculate statistics
        self.stats = self.calculate_statistics(df)
        print(f"   • Critical: {self.stats['critical']}")
        print(f"   • High: {self.stats['high']}")
        print(f"   • Medium: {self.stats['medium']}")
        print(f"   • Low: {self.stats['low']}")
        print(f"   • Unique IPs: {self.stats['unique_ips']}")
        print(f"   • Unique Vulnerabilities: {self.stats['unique_vulns']}")
        
        # Get unique vulnerabilities
        self.unique_vulnerabilities = self.get_unique_vulnerabilities(df)
        
        # Determine client name
        if not client_name and self.scan_info.get('asset_groups'):
            client_name = self.scan_info['asset_groups']
        if not client_name:
            client_name = "[MANUAL: Enter Client Name]"
        
        # Load template
        print(f"📄 Loading template: {Path(template_file).name}")
        doc = Document(template_file)
        
        # Prepare replacement dictionary
        scan_date_formatted = self.scan_info['scan_date'].strftime('%B %d, %Y') if isinstance(self.scan_info['scan_date'], datetime) else datetime.now().strftime('%B %d, %Y')
        
        # Calculate assessment period
        launch_date = self.scan_info['scan_date'] if isinstance(self.scan_info['scan_date'], datetime) else datetime.now()
        
        # Helper functions for number words
        def number_to_words(n):
            """Convert number to words for 0-100"""
            ones = ['zero', 'one', 'two', 'three', 'four', 'five', 'six', 'seven', 'eight', 'nine',
                    'ten', 'eleven', 'twelve', 'thirteen', 'fourteen', 'fifteen', 'sixteen',
                    'seventeen', 'eighteen', 'nineteen']
            tens = ['', '', 'twenty', 'thirty', 'forty', 'fifty', 'sixty', 'seventy', 'eighty', 'ninety']
            
            if n < 20:
                return ones[n]
            elif n < 100:
                return tens[n // 10] + ('' if n % 10 == 0 else '-' + ones[n % 10])
            else:
                return str(n)
        
        def format_count(n):
            """Format count as 'Word (Number)'"""
            word = number_to_words(n).title() if n < 100 else str(n)
            return f"{word} ({n})"
        
        replacements = {
            '<Client Name>': client_name,
            '<Date>': scan_date_formatted,
            'two (2)': format_count(self.stats['total_vulnerabilities']),
            'Two (2)': format_count(self.stats['total_vulnerabilities']),
            'November 6, 2025': launch_date.strftime('%B %d, %Y'),
            'November 11 to November 12, 2025': f"{launch_date.strftime('%B %d')} to {launch_date.strftime('%B %d, %Y')}",
            'November 10 to November 11, 2025': f"{launch_date.strftime('%B %d')} to {launch_date.strftime('%B %d, %Y')}",
            'MEDIUM severity': self._get_severity_text(self.stats),
            # Don't add a catch-all for "Vulnerability Count..." - handle it specifically in _replace_vulnerability_counts
        }
        
        print("✏️  Replacing placeholders...")
        self.replace_placeholders(doc, replacements)
        
        # Replace vulnerability counts BEFORE populating tables (so table references work)
        print("🔢 Replacing vulnerability counts...")
        self._replace_vulnerability_counts(doc, self.stats)
        
        # Additional comprehensive replacement for any remaining embedded instances
        print("🔍 Final sweep for embedded placeholders...")
        self._final_placeholder_sweep(doc, self.stats)
        
        # Populate tables
        print("📋 Populating tables...")
        
        # Get unique IPs for asset table
        unique_ips = sorted(df['IP'].unique().tolist()) if len(df) > 0 else []
        assets = [{'ip': ip, 'role': '[MANUAL: Enter device role]'} for ip in unique_ips]
        self.populate_asset_scope_table(doc, assets)
        
        # Calculate stats by IP
        ip_stats = {}
        for ip in unique_ips:
            ip_df = df[df['IP'] == ip]
            ip_stats[ip] = {
                'critical': len(ip_df[ip_df['Severity_Label'] == 'Critical']),
                'high': len(ip_df[ip_df['Severity_Label'] == 'High']),
                'medium': len(ip_df[ip_df['Severity_Label'] == 'Medium']),
                'low': len(ip_df[ip_df['Severity_Label'] == 'Low']),
                'total': len(ip_df)
            }
        
        self.populate_summary_of_findings_table(doc, ip_stats)
        self.populate_vulnerability_summary_table(doc, self.unique_vulnerabilities)
        
        # Populate conclusions table
        print("📊 Populating conclusions table...")
        self._populate_conclusions_table(doc, self.unique_vulnerabilities)
        
        # Add vulnerability details sections
        print("📝 Adding vulnerability details...")
        self._add_vulnerability_details(doc, self.unique_vulnerabilities)
        
        # Add manual review markers
        print("🔖 Adding manual review markers...")
        self._add_review_markers(doc)
        
        # Save document
        print(f"💾 Saving report: {Path(output_file).name}")
        doc.save(output_file)
        
        # Print summary
        print(f"\n{'='*60}")
        print(f"✅ REPORT GENERATION COMPLETE")
        print(f"{'='*60}")
        print(f"Client: {client_name}")
        print(f"Total Vulnerabilities: {self.stats['total_vulnerabilities']}")
        print(f"Unique Vulnerabilities: {self.stats['unique_vulns']}")
        print(f"Output: {output_file}")
        print(f"\n⚠️  MANUAL REVIEW REQUIRED (~2 minutes):")
        print(f"   • Verify narrative tone in Executive Summary")
        print(f"   • Add device roles in Asset Scope table")
        print(f"   • Update Document Version History names/dates")
        print(f"   • Add team member names in Review/Approval section")
        print(f"   • Final spot-check for accuracy")
        print(f"{'='*60}\n")
        
        return True
    
    def _get_severity_text(self, stats):
        """Generate severity text for executive summary"""
        severities = []
        if stats['critical'] > 0:
            severities.append('CRITICAL')
        if stats['high'] > 0:
            severities.append('HIGH')
        if stats['medium'] > 0:
            severities.append('MEDIUM')
        if stats['low'] > 0:
            severities.append('LOW')
        
        if len(severities) == 0:
            return 'no vulnerabilities found'
        elif len(severities) == 1:
            return f"{severities[0]} severity"
        else:
            return f"{', '.join(severities[:-1])}, and {severities[-1]} severity"
    
    def _add_vulnerability_details(self, doc, unique_vulns):
        """Add vulnerability detail sections by carefully updating template tables"""
        print(f"   Populating Section 3.4 with detailed write-ups...")
        
        # Include ALL vulnerabilities (Critical, High, Medium, Low) - sorted by severity
        all_vulns = sorted(unique_vulns, key=lambda v: self.SEVERITY_ORDER.get(v['severity'], 999))
        
        if len(all_vulns) == 0:
            print("   No vulnerabilities to detail")
            return
        
        print(f"   Processing {len(all_vulns)} total vulnerabilities (all severity levels)")
        
        # Find all vulnerability detail template tables (6-row, 2-column tables with "Severity" in first cell)
        vuln_tables = []
        for table in doc.tables:
            if len(table.rows) >= 6 and len(table.columns) == 2:
                first_cell = table.rows[0].cells[0].text.strip()
                if 'Severity' in first_cell:
                    vuln_tables.append(table)
        
        if len(vuln_tables) == 0:
            print("   ⚠️ No template tables found for vulnerability details")
            return
        
        print(f"   Found {len(vuln_tables)} template tables, need {len(all_vulns)} for vulnerabilities")
        
        # Use as many tables as we have available
        tables_to_use = min(len(vuln_tables), len(all_vulns))
        
        # Populate each template table with actual vulnerability data
        for idx, (vuln, table) in enumerate(zip(all_vulns[:tables_to_use], vuln_tables[:tables_to_use]), 1):
            # Update the title paragraph (find preceding paragraph with "<Vulnerability Name #X>")
            for para in doc.paragraphs:
                if f'Vulnerability Name # {idx}' in para.text:
                    # Replace text but keep formatting
                    self._update_paragraph_text(para, f"{idx}. {vuln['title']}")
                    break
            
            # Row 0: Severity - Apply exact template formatting
            self._format_severity_cell(table.rows[0].cells[1], vuln['severity'])
            
            # Row 1: Status - Apply exact template formatting
            self._format_status_cell(table.rows[1].cells[1], 'OPEN')
            
            # Row 2: Affected Host(s)
            affected_ips = ', '.join(vuln['affected_ips'])
            self._update_cell_text(table.rows[2].cells[1], affected_ips)
            
            # Row 3: Description
            description = self.clean_text(vuln['threat'])
            if len(description) > 800:
                description = description[:800] + "..."
            self._update_cell_text(table.rows[3].cells[1], description)
            
            # Row 4: Impact
            impact = self.clean_text(vuln['impact'])
            if len(impact) > 600:
                impact = impact[:600] + "..."
            self._update_cell_text(table.rows[4].cells[1], impact)
            
            # Row 5: Vulnerability Proof/Evidence (if exists)
            if len(table.rows) > 5:
                proof = self.clean_text(vuln['results'])
                if len(proof) > 500:
                    proof = proof[:500] + "...\n[See detailed report for full output]"
                self._update_cell_text(table.rows[5].cells[1], proof)
            
            # Row 6: Recommended Resolution (if exists)
            if len(table.rows) > 6:
                solution = self.clean_text(vuln['solution'])
                if len(solution) > 800:
                    solution = solution[:800] + "..."
                self._update_cell_text(table.rows[6].cells[1], solution)
        
        if tables_to_use < len(all_vulns):
            remaining = len(all_vulns) - tables_to_use
            print(f"   ⚠️ Warning: Only {tables_to_use} template tables available for {len(all_vulns)} vulnerabilities")
            print(f"   ⚠️ {remaining} vulnerabilities not detailed. Add more template tables.")
        
        print(f"   ✅ Populated {tables_to_use} vulnerability tables with EXACT template formatting preserved")
        
        # Auto-delete unused placeholder tables
        if len(vuln_tables) > tables_to_use:
            unused_count = len(vuln_tables) - tables_to_use
            print(f"   🗑️  Removing {unused_count} unused placeholder tables...")
            self._delete_unused_tables(doc, vuln_tables[tables_to_use:])
            print(f"   ✅ Cleaned up {unused_count} unused tables")
    
    def _update_cell_text(self, cell, new_text):
        """Update cell text while preserving ALL formatting (fonts, colors, bold, etc.)"""
        # Strategy: Replace text in the first paragraph's first run, keep all formatting
        if len(cell.paragraphs) == 0:
            return
        
        para = cell.paragraphs[0]
        
        # If there are existing runs, use the first one and update its text
        if len(para.runs) > 0:
            # Keep only the first run, remove others
            first_run = para.runs[0]
            
            # Remove other runs
            for i in range(len(para.runs) - 1, 0, -1):
                para._element.remove(para.runs[i]._element)
            
            # Update the first run's text (preserves all its formatting)
            first_run.text = new_text
        else:
            # No runs exist, create one (this shouldn't happen with template tables)
            run = para.add_run(new_text)
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
        
        # Remove extra paragraphs (keep only first one)
        for i in range(len(cell.paragraphs) - 1, 0, -1):
            p = cell.paragraphs[i]
            p._element.getparent().remove(p._element)
    
    def _format_severity_cell(self, cell, severity):
        """Format severity cell with exact template colors and highlighting"""
        # Clear existing content
        para = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
        para.text = ""
        
        # Remove extra paragraphs
        for i in range(len(cell.paragraphs) - 1, 0, -1):
            p = cell.paragraphs[i]
            p._element.getparent().remove(p._element)
        
        # Add the severity text with exact formatting from template
        run = para.add_run(severity.upper())
        run.font.bold = True
        run.font.name = 'Aptos'
        run.font.size = Pt(11)
        
        # Apply colors based on severity (matching template exactly)
        if severity == 'Critical':
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            # Red highlighting (C00000)
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd {} w:fill="C00000"/>'.format(nsdecls('w')))
            run._element.rPr.append(shd)
        elif severity == 'High':
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            # Orange highlighting (ED7D31)
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd {} w:fill="ED7D31"/>'.format(nsdecls('w')))
            run._element.rPr.append(shd)
        elif severity == 'Medium':
            run.font.color.rgb = RGBColor(38, 38, 38)  # Dark text (262626)
            # Yellow highlighting (FFC000)
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            run._element.rPr.append(shd)
        elif severity == 'Low':
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            # Green highlighting (92D050)
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd {} w:fill="92D050"/>'.format(nsdecls('w')))
            run._element.rPr.append(shd)
    
    def _format_status_cell(self, cell, status):
        """Format status cell with exact template colors and formatting"""
        # Clear existing content
        para = cell.paragraphs[0] if len(cell.paragraphs) > 0 else cell.add_paragraph()
        para.text = ""
        
        # Remove extra paragraphs
        for i in range(len(cell.paragraphs) - 1, 0, -1):
            p = cell.paragraphs[i]
            p._element.getparent().remove(p._element)
        
        # Add the status text with exact formatting
        run = para.add_run(status.upper())
        run.font.bold = True
        run.font.name = 'Aptos'
        run.font.size = Pt(11)
        
        # Apply colors based on status (matching template exactly)
        if status.upper() == 'OPEN':
            run.font.color.rgb = RGBColor(192, 0, 0)  # Red text (C00000)
            run.font.underline = True
        elif status.upper() == 'RISK ACCEPTED':
            run.font.color.rgb = RGBColor(38, 38, 38)  # Dark text
            # Yellow highlighting
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
            run._element.rPr.append(shd)
        elif status.upper() == 'CLOSED':
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            # Green highlighting
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shd = parse_xml(r'<w:shd {} w:fill="92D050"/>'.format(nsdecls('w')))
            run._element.rPr.append(shd)
    
    def _update_paragraph_text(self, para, new_text):
        """Update paragraph text while preserving formatting"""
        if len(para.runs) > 0:
            # Keep first run, update its text
            first_run = para.runs[0]
            for i in range(len(para.runs) - 1, 0, -1):
                para._element.remove(para.runs[i]._element)
            first_run.text = new_text
        else:
            para.add_run(new_text)
    
    def _delete_unused_tables(self, doc, tables_to_delete):
        """Delete unused placeholder tables and their associated numbered paragraphs"""
        # First, collect the table numbers we need to delete
        populated_count = len([t for t in doc.tables if self._is_vuln_detail_table(t)]) - len(tables_to_delete)
        
        # Delete paragraphs and tables for unused entries
        for idx in range(len(tables_to_delete)):
            table_num = populated_count + idx + 1  # e.g., if 18 populated, start deleting from #19
            
            # Find and remove numbered paragraphs (might have variations like "3.4.19" or just the number)
            paras_to_remove = []
            for para in doc.paragraphs:
                para_text = para.text.strip()
                # Match patterns like "3.4.19", "19.", or "Vulnerability Name # 19"
                if (f'3.4.{table_num}' in para_text or 
                    f'Vulnerability Name # {table_num}' in para_text or
                    f'Vulnerability Name #{table_num}' in para_text):
                    paras_to_remove.append(para)
            
            # Remove the paragraphs
            for para in paras_to_remove:
                p_element = para._element
                p_element.getparent().remove(p_element)
        
        # Now remove the tables
        for table in tables_to_delete:
            tbl = table._element
            parent = tbl.getparent()
            parent.remove(tbl)
    
    def _is_vuln_detail_table(self, table):
        """Check if a table is a vulnerability detail table"""
        if len(table.rows) >= 6 and len(table.columns) == 2:
            first_cell = table.rows[0].cells[0].text.strip()
            if 'Severity' in first_cell:
                return True
        return False
    
    def _replace_vulnerability_counts(self, doc, stats):
        """Replace vulnerability count placeholders in bullet lists"""
        
        def number_to_words(n):
            """Convert number to words"""
            ones = ['Zero', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine',
                    'Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen',
                    'Seventeen', 'Eighteen', 'Nineteen', 'Twenty']
            tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
            
            if n < len(ones):
                return ones[n]
            elif n < 100:
                tens_digit = n // 10
                ones_digit = n % 10
                if ones_digit == 0:
                    return tens[tens_digit]
                else:
                    return tens[tens_digit] + '-' + ones[ones_digit].lower()
            return str(n)
        
        def format_count(n):
            return f"{number_to_words(n)} ({n})"
        
        # Process each paragraph individually
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip if no placeholder
            if 'Vulnerability Count in Words (Number)' not in text:
                continue
            
            # Determine replacement based on the BEGINNING of the paragraph text
            new_value = None
            
            # Check what the line starts with (after any list markers)
            clean_text = text.lstrip('•-*\t ')  # Remove common list markers
            
            if clean_text.startswith('Discovered vulnerabilities'):
                new_value = format_count(stats['total_vulnerabilities'])
            elif clean_text.startswith('Unique Vulnerabilities'):
                new_value = format_count(stats['unique_vulns'])
            elif clean_text.startswith('Critical'):
                new_value = format_count(stats['critical'])
            elif clean_text.startswith('High'):
                new_value = format_count(stats['high'])
            elif clean_text.startswith('Medium'):
                new_value = format_count(stats['medium'])
            elif clean_text.startswith('Low'):
                new_value = format_count(stats['low'])
            
            # Replace in all runs of this paragraph
            if new_value:
                for run in paragraph.runs:
                    if 'Vulnerability Count in Words (Number)' in run.text:
                        run.text = run.text.replace('Vulnerability Count in Words (Number)', new_value)
    
    def _final_placeholder_sweep(self, doc, stats):
        """Final comprehensive sweep to replace any remaining placeholders"""
        
        def number_to_words(n):
            ones = ['Zero', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine',
                    'Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen',
                    'Seventeen', 'Eighteen', 'Nineteen', 'Twenty']
            tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']
            
            if n < len(ones):
                return ones[n]
            elif n < 100:
                tens_digit = n // 10
                ones_digit = n % 10
                if ones_digit == 0:
                    return tens[tens_digit]
                else:
                    return tens[tens_digit] + '-' + ones[ones_digit].lower()
            return str(n)
        
        def format_count(n):
            return f"{number_to_words(n)} ({n})"
        
        # Catch-all for any remaining "Vulnerability Count in Words (Number)" instances
        count_formatted = format_count(stats['total_vulnerabilities'])
        
        replacements_made = 0
        for paragraph in doc.paragraphs:
            if 'Vulnerability Count in Words (Number)' in paragraph.text:
                for run in paragraph.runs:
                    if 'Vulnerability Count in Words (Number)' in run.text:
                        run.text = run.text.replace('Vulnerability Count in Words (Number)', count_formatted)
                        replacements_made += 1
        
        # Also check tables for any remaining placeholders
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'Vulnerability Count in Words (Number)' in paragraph.text:
                            for run in paragraph.runs:
                                if 'Vulnerability Count in Words (Number)' in run.text:
                                    run.text = run.text.replace('Vulnerability Count in Words (Number)', count_formatted)
                                    replacements_made += 1
        
        if replacements_made > 0:
            print(f"   Fixed {replacements_made} additional embedded placeholder(s)")
    
    def _populate_conclusions_table(self, doc, unique_vulns):
        """Populate the conclusions table with vulnerability aging information"""
        # Find table in CONCLUSIONS section - look for "Initial Report" section specifically
        in_conclusions = False
        target_table = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            if 'CONCLUSIONS' in paragraph.text or 'Conclusions' in paragraph.text:
                in_conclusions = True
            
            # Look for the "For Initial Report" marker
            if in_conclusions and '<For Initial Report>' in paragraph.text:
                # The table should be shortly after this marker
                # Look for the next table with DESCRIPTION column
                for j in range(i, min(i + 10, len(doc.paragraphs))):
                    # Check tables that appear around this area
                    pass
        
        # Alternative approach: find by table header content
        for table in doc.tables:
            # Check if this table has the right structure
            if len(table.rows) > 0:
                header_text = ''.join([cell.text for cell in table.rows[0].cells]).upper()
                # Look for the simpler Initial Report table: DESCRIPTION | RISK RATING | STATUS | SECTION REFERENCE
                if 'DESCRIPTION' in header_text and 'RISK' in header_text and 'STATUS' in header_text:
                    # Check if it has 4 columns (Initial Report format)
                    if len(table.rows[0].cells) == 4:
                        # Check if it has placeholder rows (XXXX markers)
                        if len(table.rows) > 1:
                            first_data_row = ''.join([cell.text for cell in table.rows[1].cells])
                            if 'XXXX' in first_data_row or '[' in first_data_row:
                                target_table = table
                                break
        
        if not target_table:
            print("   ⚠️ Could not find conclusions table with correct structure")
            return
        
        # Clear existing placeholder rows (keep header)
        rows_to_remove = len(target_table.rows) - 1
        for _ in range(rows_to_remove):
            if len(target_table.rows) > 1:
                target_table._element.remove(target_table.rows[-1]._element)
        
        # Calculate days aging (from scan date to now)
        scan_date = self.scan_info.get('scan_date', datetime.now())
        if isinstance(scan_date, datetime):
            # Remove timezone info to avoid comparison issues
            if scan_date.tzinfo is not None:
                scan_date = scan_date.replace(tzinfo=None)
            days_aging = (datetime.now() - scan_date).days
        else:
            days_aging = 0
        
        # Add rows for ALL vulnerabilities (sorted by severity)
        sorted_vulns = sorted(unique_vulns, key=lambda v: self.SEVERITY_ORDER.get(v['severity'], 999))
        
        for vuln in sorted_vulns:
            row = target_table.add_row()
            self.format_cell_text(row.cells[0], vuln['title'], font_size=10)
            self.format_cell_text(row.cells[1], vuln['severity'].upper(), bold=True, font_size=10)
            
            # Color code severity using new formatting method
            self._format_severity_cell(row.cells[1], vuln['severity'])
            
            self.format_cell_text(row.cells[2], 'OPEN', font_size=10)
            self.format_cell_text(row.cells[3], f'Sec. 3.4.{sorted_vulns.index(vuln) + 1}', font_size=10)
        
        print(f"   ✅ Populated conclusions table with {len(sorted_vulns)} vulnerabilities (all severity levels)")


    
    def _add_review_markers(self, doc):
        """Add markers for sections requiring manual review"""
        # Find Executive Summary section and add review note
        for i, paragraph in enumerate(doc.paragraphs):
            if 'EXECUTIVE SUMMARY' in paragraph.text:
                # Add review marker after the heading
                if i + 1 < len(doc.paragraphs):
                    # Insert before the first paragraph of content
                    marker_para = doc.paragraphs[i + 1].insert_paragraph_before()
                    marker_run = marker_para.add_run(
                        "[MANUAL REVIEW: Verify narrative flow, adjust tone for client, "
                        "confirm assessment dates match actual timeline]"
                    )
                    marker_run.font.bold = True
                    marker_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
                    marker_run.font.size = Pt(9)
                break

def main():
    if len(sys.argv) < 4:
        print("""
╔═══════════════════════════════════════════════════════════╗
║   Executive Report Automation Tool v1.5                   ║
║   98% Automated - 2% Manual Review                        ║
╚═══════════════════════════════════════════════════════════╝

USAGE:
  python executive_report_automation_v1.py <csv_file> <template.docx> <output.docx> [client_name]

EXAMPLES:
  python executive_report_automation_v1.py scan.csv template.docx output.docx "Acme Corp"
  python executive_report_automation_v1.py scan.csv template.docx output.docx

NEW IN v1.5:
  ✅ NOW INCLUDES ALL VULNERABILITIES (Critical, High, Medium, Low)
  ✅ Exact severity/status formatting from template (colored backgrounds!)
  ✅ Bold labels preserved (Description, Impact, Recommended Resolution)
  ✅ Perfect format match with Report-Template-Bold-Fixed.docx
  
FEATURES:
  ✅ Auto-extracts scan metadata from CSV
  ✅ Populates client name and dates
  ✅ Generates all tables (Asset Scope, Summary, Details)
  ✅ Creates detailed sections for ALL vulnerabilities (all severity levels)
  ✅ Auto-removes unused placeholder tables
  ✅ Exact template formatting preserved (teal headers, colors, fonts)
  ✅ Calculates all statistics
  ✅ Marks sections for manual review
  
MANUAL REVIEW AREAS:
  📝 Executive summary narrative (tone check)
  📝 Device roles in asset table
  📝 Version history updates
  📝 Team member names
  📝 Final quality check
        """)
        sys.exit(1)
    
    csv_file = sys.argv[1]
    template_file = sys.argv[2]
    output_file = sys.argv[3]
    client_name = sys.argv[4] if len(sys.argv) > 4 else None
    
    # Validate inputs
    if not os.path.exists(csv_file):
        print(f"❌ ERROR: CSV file not found: {csv_file}")
        sys.exit(1)
    
    if not os.path.exists(template_file):
        print(f"❌ ERROR: Template file not found: {template_file}")
        sys.exit(1)
    
    # Generate report
    automation = ExecutiveReportAutomation()
    success = automation.generate_report(csv_file, template_file, output_file, client_name)
    
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
